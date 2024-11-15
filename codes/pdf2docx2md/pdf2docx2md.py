"""
!pip install pdf2docx
!pip install python-docx



# To run this script from the command line, use:
# python pdf2docx2MD.py --input_pdf <path_to_pdf_file> --output_dir <path_to_save_output_files>
#
# Example:
# python pdf2docx2md.py --input_pdf "./input.pdf" --output_dir "./output"
# python pdf2docx2md.py --input_pdf "./08_law_main.pdf" --output_dir "./output"
#
# This script converts a PDF file to a DOCX file and then converts the DOCX file to Markdown (MD) format.
# - The first three lines of the DOCX file are removed after conversion.
# - The Markdown file is saved in the specified output directory.
"""

import os
import shutil
import argparse
import sys
from pdf2docx import Converter
import docx

# Add the project's root directory to sys.path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

# Argument group
parser = argparse.ArgumentParser(description="Convert PDF to DOCX and then to Markdown.")
g = parser.add_argument_group("Arguments")

g.add_argument("--input_pdf", type=str, required=True, help="Path to the input PDF file.")
g.add_argument("--output_dir", type=str, required=True, help="Path where the output files will be saved.")

args = parser.parse_args()

def pdf2docx(pdf_file, docx_output):
    """Convert PDF to DOCX."""
    cv = Converter(pdf_file)
    cv.convert(docx_output, start=0, end=None, image_quality=0)  # 조정된 이미지 품질 설정
    cv.close()

def delete_lines(docx_file, line_num=3):
    """Delete the first three lines from a DOCX file."""
    doc = docx.Document(docx_file)
    for i in range(line_num):
        if len(doc.paragraphs) > i:
            doc.paragraphs[i].clear()
    
    # Remove empty paragraphs
    while len(doc.paragraphs) > 0 and doc.paragraphs[0].text == '':
        doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._p)
    
    doc.save(docx_file)

def docx2md(md_output, docx_file):
    """Convert DOCX to Markdown using Pandoc."""
    md_file = md_output
    media_dir = os.path.join(os.path.dirname(md_file), 'media')

    # Use pandoc to convert docx to markdown
    cmd = f'pandoc --extract-media=./ --markdown-headings=atx --wrap=none --toc -f docx -t markdown_strict "{docx_file}" -o "{md_file}"'
    os.system(cmd)

    # Move media files if any exist
    if os.path.exists(media_dir):
        shutil.move(media_dir, os.path.dirname(md_file))  # Move media files

def pdf2md(pdf_file, output_dir):
    """Convert PDF to Markdown via DOCX."""
    # Extract base name from PDF file path
    base_name = os.path.splitext(os.path.basename(pdf_file))[0]

    # Step 1: Convert PDF to DOCX
    docx_file = os.path.join(output_dir, base_name + '.docx')
    pdf2docx(pdf_file, docx_file)

    # Step 2: Remove the first three lines from the DOCX file
    delete_lines(docx_file)

    # Step 3: Convert the cleaned DOCX to Markdown
    md_file = os.path.join(output_dir, base_name + '.md')
    docx2md(md_file, docx_file)

    print(f"Conversion from PDF to Markdown for {base_name} completed! \nConverted Files saved in {output_dir}")

def main(args):
    # Ensure output directory exists
    os.makedirs(args.output_dir, exist_ok=True)

    # Call the conversion function with provided arguments
    pdf2md(args.input_pdf, args.output_dir)

if __name__ == "__main__":
    exit(main(parser.parse_args()))
